from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_from_directory
import os
from dotenv import load_dotenv
from wordcloud import WordCloud
from services.azure_auth import AzureAuth
from services.azure_services import LocalNLPService, AzureAIService, AzureBlobStorageService, AzureTableAuditService
from comparison.semantic_analysis import SemanticAnalyzer
from comparison.sentiment_risk import SentimentRiskClassifier
from comparison.tone_shift import ToneShiftAnalyzer
from comparison.diff_view import DiffView
from comparison.insights import InsightsGenerator
from comparison.multilingual import MultilingualService
from utils.metrics import get_all_metrics
from utils.helpers import log_error, is_supported_filetype
from models import Document, AnalysisResult
import io
import base64

load_dotenv()
UPLOAD_FOLDER = 'uploads'
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your_default_secret_key')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize Azure services and analyzers
azure_auth = AzureAuth(
    os.getenv('AZURE_CLIENT_ID'),
    os.getenv('AZURE_CLIENT_SECRET'),
    os.getenv('AZURE_TENANT_ID')
)
azure_ai_service = AzureAIService(
    os.getenv('AZURE_AI_ENDPOINT'),
    os.getenv('AZURE_AI_API_KEY'),
    region=os.getenv('AZURE_AI_REGION'),
    deployment_name=os.getenv('AZURE_AI_DEPLOYMENT', 'gpt-4')
)
multilingual_service = MultilingualService(
    subscription_key=os.getenv('AZURE_TRANSLATOR_KEY'),
    endpoint=os.getenv('AZURE_TRANSLATOR_ENDPOINT'),
    region=os.getenv('AZURE_TRANSLATOR_REGION'),
    azure_ai_service=azure_ai_service
)
local_nlp_service = LocalNLPService()
semantic_analyzer = SemanticAnalyzer(local_nlp_service, multilingual_service, azure_ai_service)
sentiment_classifier = SentimentRiskClassifier(local_nlp_service, multilingual_service, azure_ai_service)
tone_shift_analyzer = ToneShiftAnalyzer(local_nlp_service, multilingual_service, azure_ai_service)
diff_view = DiffView(multilingual_service, azure_ai_service)
insights_generator = InsightsGenerator(semantic_analyzer, sentiment_classifier, tone_shift_analyzer, diff_view, azure_ai_service)
AZURE_BLOB_CONNECTION_STRING = os.getenv('AZURE_BLOB_CONNECTION_STRING')
AZURE_BLOB_CONTAINER = os.getenv('AZURE_BLOB_CONTAINER', 'documents')
AZURE_TABLE_CONNECTION_STRING = os.getenv('AZURE_TABLE_CONNECTION_STRING')
AZURE_TABLE_NAME = os.getenv('AZURE_TABLE_NAME', 'AuditLog')
blob_service = AzureBlobStorageService(AZURE_BLOB_CONNECTION_STRING, AZURE_BLOB_CONTAINER)
audit_service = AzureTableAuditService(AZURE_TABLE_CONNECTION_STRING, AZURE_TABLE_NAME)

def generate_wordcloud(text):
    if not text.strip():
        return ""
    wc = WordCloud(width=800, height=400, background_color='white').generate(text)
    img_io = io.BytesIO()
    wc.to_image().save(img_io, format='PNG')
    img_io.seek(0)
    return base64.b64encode(img_io.read()).decode('utf-8')

def allowed_file(filename):
    return is_supported_filetype(filename)

def read_file_content(filepath):
    ext = filepath.rsplit('.', 1)[-1].lower()
    if ext in ['txt', 'md']:
        with open(filepath, encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == 'docx':
        from docx import Document as DocxDocument
        doc = DocxDocument(filepath)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == 'pdf':
        try:
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            return "PyPDF2 not installed. Please install it to support PDF files."
    elif ext == 'csv':
        import pandas as pd
        df = pd.read_csv(filepath)
        return df.to_string(index=False)
    elif ext == 'xlsx':
        import pandas as pd
        df = pd.read_excel(filepath)
        return df.to_string(index=False)
    return ""

@app.route('/health')
def health():
    return {
        "nlp_service": "LocalNLPService",
        "semantic_analyzer": semantic_analyzer.health_check(),
        "sentiment_classifier": sentiment_classifier.health_check(),
        "tone_shift_analyzer": tone_shift_analyzer.health_check()
    }

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        files = [request.files.get('document1'), request.files.get('document2')]
        docs, filenames = [], []
        for file in files:
            if file and allowed_file(file.filename):
                filename = file.filename
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                try:
                    content = read_file_content(filepath)
                except Exception as e:
                    log_error(f"Failed to read file {filename}", exc=e)
                    content = ""
                docs.append(content)
                filenames.append(filename)
            else:
                docs.append("")
                filenames.append("")
        if not docs[0] or not docs[1]:
            flash("Please upload two valid documents.", "danger")
            return redirect(url_for('index'))
        return redirect(url_for('compare', doc1_name=filenames[0], doc2_name=filenames[1]))
    health_status = {
        "nlp_service": "LocalNLPService",
        "semantic": semantic_analyzer.health_check(),
        "sentiment": sentiment_classifier.health_check(),
        "tone": tone_shift_analyzer.health_check()
    }
    return render_template('index.html', health_status=health_status)

@app.route('/compare')
def compare():
    doc1_name = request.args.get('doc1_name')
    doc2_name = request.args.get('doc2_name')
    doc1_path = os.path.join(app.config['UPLOAD_FOLDER'], doc1_name)
    doc2_path = os.path.join(app.config['UPLOAD_FOLDER'], doc2_name)
    try:
        doc1_content = read_file_content(doc1_path)
        doc2_content = read_file_content(doc2_path)
        wordcloud1 = generate_wordcloud(doc1_content)
        wordcloud2 = generate_wordcloud(doc2_content)
        doc1 = Document(title=doc1_name, content=doc1_content)
        doc2 = Document(title=doc2_name, content=doc2_content)
    except Exception as e:
        log_error(f"Error reading files {doc1_name} or {doc2_name}", exc=e)
        flash("Error reading uploaded files.", "danger")
        return redirect(url_for('index'))
    try:
        insights = insights_generator.get_document_insights(doc1.content, doc2.content)
        metrics = get_all_metrics(doc1.content, doc2.content)
        analysis_result = AnalysisResult(
            document=doc1,
            semantic_analysis=insights.get("semantic_report", {}),
            sentiment=metrics.get("sentiment_doc1", ""),
            tone_shifts=[insights.get("tone_report", "")],
            diff=insights.get("diff_stats", {}),
            compliance_flags=insights.get("compliance_flags_doc1", []),
            pii_entities=insights.get("pii_doc1", []),
            ai_insights=insights,
            metrics=metrics
        )
    except Exception as e:
        log_error("Error generating insights or metrics", exc=e)
        flash("Error processing documents.", "danger")
        return redirect(url_for('index'))
    try:
        store_document_and_result(doc1_name, doc2_name, doc1.content, doc2.content, analysis_result.to_dict())
    except Exception as e:
        log_error("Error storing documents/results", exc=e)
        flash("Error storing results in database.", "warning")
    return render_template('compare.html', doc1_name=doc1_name, doc2_name=doc2_name, insights=insights, metrics=metrics, wordcloud1=wordcloud1, wordcloud2=wordcloud2)

@app.route('/download/<filename>')
def download(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

@app.route('/advanced', methods=['GET', 'POST'])
def advanced():
    doc1 = doc2 = ""
    doc1_name = doc2_name = ""
    result = {}
    if request.method == 'POST':
        files = [request.files.get('document1'), request.files.get('document2')]
        docs = []
        for file in files:
            if file and allowed_file(file.filename):
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(filepath)
                with open(filepath, encoding='utf-8', errors='ignore') as f:
                    docs.append(f.read())
            else:
                docs.append("")
        doc1, doc2 = docs
        doc1_name = files[0].filename if files[0] else ""
        doc2_name = files[1].filename if files[1] else ""
        result = {
            "diff_view": diff_view.generate_diff_view(doc1, doc2),
            "side_by_side_html": diff_view.generate_side_by_side_html(doc1, doc2),
            "inline_diff_html": diff_view.generate_inline_diff(doc1, doc2),
            "changed_lines": diff_view.get_changed_lines(doc1, doc2),
            "diff_stats": diff_view.get_diff_stats(doc1, doc2),
            "similarity_score": diff_view.get_similarity_score(doc1, doc2),
            "summarize_diff_changes": diff_view.summarize_diff_changes(doc1, doc2),
            "contextual_diff": diff_view.get_contextual_diff(doc1, doc2),

            "sentiment_details_doc1": sentiment_classifier.get_sentiment_details(doc1),
            "sentiment_details_doc2": sentiment_classifier.get_sentiment_details(doc2),
            "priority_action_report_doc1": sentiment_classifier.get_priority_action_report(doc1),
            "priority_action_report_doc2": sentiment_classifier.get_priority_action_report(doc2),
            "sentiment_time_series": sentiment_classifier.get_sentiment_time_series([doc1, doc2], ["Doc1", "Doc2"]),
            "sentiment_change_points": sentiment_classifier.get_sentiment_change_points([doc1, doc2]),

            "tone_shift_details": tone_shift_analyzer.get_tone_shift_details(doc1, doc2),
            "tone_distribution": tone_shift_analyzer.get_tone_distribution([doc1, doc2]),
            "compliance_flags_doc1": tone_shift_analyzer.get_compliance_flags(doc1),
            "compliance_flags_doc2": tone_shift_analyzer.get_compliance_flags(doc2),
            "highlight_important_tone_points_doc1": tone_shift_analyzer.highlight_important_tone_points(doc1),
            "highlight_important_tone_points_doc2": tone_shift_analyzer.highlight_important_tone_points(doc2),
            "tone_trend": tone_shift_analyzer.get_tone_trend([doc1, doc2]),
            
            "semantic_analysis": semantic_analyzer.analyze_semantics(doc1, doc2),
            "extract_entities_doc1": semantic_analyzer.extract_entities(doc1),
            "extract_entities_doc2": semantic_analyzer.extract_entities(doc2),
            "compare_entities": semantic_analyzer.compare_entities(doc1, doc2),

            "detected_language_doc1": multilingual_service.detect_language(doc1),
            "detected_language_doc2": multilingual_service.detect_language(doc2),
            "summarize_text_doc1": multilingual_service.summarize_text(doc1),
            "summarize_text_doc2": multilingual_service.summarize_text(doc2),
            "detect_pii_doc1": multilingual_service.detect_pii(doc1),
            "detect_pii_doc2": multilingual_service.detect_pii(doc2),
            "extract_topics_doc1": multilingual_service.extract_topics(doc1),
            "extract_topics_doc2": multilingual_service.extract_topics(doc2),
        }
    return render_template('advanced.html', doc1_name=doc1_name, doc2_name=doc2_name, result=result)

@app.route('/audit')
def audit():
    user_id = request.args.get('user_id', None)
    logs = audit_service.get_audit_logs(user_id)
    return render_template('audit.html', logs=logs)

@app.route('/save_to_db', methods=['POST'])
def save_to_db():
    data = request.get_json()
    doc1_name = data.get('doc1_name')
    doc2_name = data.get('doc2_name')
    doc1_content = data.get('doc1_content')
    doc2_content = data.get('doc2_content')
    result_summary = data.get('result_summary')
    user_id = data.get('user_id', None)
    try:
        store_document_and_result(doc1_name, doc2_name, doc1_content, doc2_content, result_summary, user_id)
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

def store_document_and_result(doc1_name, doc2_name, doc1_content, doc2_content, result_summary, user_id=None):
    blob_service.upload_text(f"{doc1_name}", doc1_content)
    blob_service.upload_text(f"{doc2_name}", doc2_content)
    result_blob_name = f"result_{doc1_name}_vs_{doc2_name}.json"
    import json
    blob_service.upload_text(result_blob_name, json.dumps(result_summary, indent=2))
    audit_service.log_audit(
        user_id=user_id,
        action="compare",
        doc1_name=doc1_name,
        doc2_name=doc2_name,
        result_summary=str(result_summary)[:500],
        status="Success"
    )

if __name__ == "__main__":
    app.run(debug=True, port=8080)